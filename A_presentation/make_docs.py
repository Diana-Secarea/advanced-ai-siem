#!/usr/bin/env python3
"""Generate two Word documents in A_presentation/:

  1. AI_Threat_Engine_Extended_Report.docx
     — abstract + system-hardening (mobile & AI apps) + reliability + pentest
       plan + guardrails + more-agentic + go-to-market.
  2. Thesis_Improvement_and_Admissions_Guide.docx
     — how to strengthen the thesis for ETH Zurich / EPFL / Europe / Australia,
       better academic titles, program positioning, application narrative.

Run:  ./ai_threat_engine_starter/venv/bin/python3 A_presentation/make_docs.py
Requires: python-docx (installed in the project venv).
Touches ONLY A_presentation/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# ---- palette (deep navy / academic) ----
NAVY   = RGBColor(0x14, 0x2A, 0x4A)
CYAN   = RGBColor(0x10, 0x6B, 0x8C)
VIOLET = RGBColor(0x5B, 0x3E, 0x9C)
GREEN  = RGBColor(0x1E, 0x7A, 0x5E)
AMBER  = RGBColor(0x9A, 0x6B, 0x00)
RED    = RGBColor(0xA3, 0x2A, 0x3C)
GREY   = RGBColor(0x55, 0x5B, 0x66)
BLACK  = RGBColor(0x1A, 0x1D, 0x24)


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------
def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.4
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for lvl, sz, col in [("Heading 1", 17, NAVY), ("Heading 2", 14, CYAN), ("Heading 3", 12, VIOLET)]:
        s = doc.styles[lvl]
        s.font.name = "Georgia"; s.font.size = Pt(sz); s.font.bold = True; s.font.color.rgb = col
        s.paragraph_format.space_before = Pt(14); s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text, italic=False, color=None, size=None, align_just=True):
    para = doc.add_paragraph()
    if align_just:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text)
    r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return para


def bullet(doc, text, bold_lead=None):
    para = doc.add_paragraph(style="List Bullet")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        r = para.add_run(bold_lead); r.bold = True
        para.add_run(" " + text)
    else:
        para.add_run(text)
    return para


def numbered(doc, text, bold_lead=None):
    para = doc.add_paragraph(style="List Number")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        r = para.add_run(bold_lead); r.bold = True
        para.add_run(" " + text)
    else:
        para.add_run(text)
    return para


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def page_break(doc):
    doc.add_page_break()


def add_footer_pagenum(doc):
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
    para._p.append(fld1)


def add_toc(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose “Update Field” to build the Table of Contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep); run._r.append(t); run._r.append(fldEnd)


def cover(doc, title, subtitle, tagline):
    doc.add_paragraph().add_run().add_break()
    doc.add_paragraph().add_run().add_break()
    pk = doc.add_paragraph(); pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pk.add_run("BUCHAREST UNIVERSITY OF ECONOMIC STUDIES")
    r.font.size = Pt(11); r.font.color.rgb = GREY; r.bold = True
    pk2 = doc.add_paragraph(); pk2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = pk2.add_run("Faculty of Cybernetics, Statistics and Economic Informatics")
    r2.font.size = Pt(10); r2.font.color.rgb = GREY
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title); tr.font.size = Pt(24); tr.bold = True; tr.font.color.rgb = NAVY
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle); sr.font.size = Pt(15); sr.font.color.rgb = CYAN; sr.italic = True
    doc.add_paragraph()
    gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = gp.add_run(tagline); gr.font.size = Pt(11); gr.font.color.rgb = GREY
    for _ in range(6): doc.add_paragraph()
    au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = au.add_run("Secărea Diana Maria"); ar.font.size = Pt(13); ar.bold = True; ar.font.color.rgb = BLACK
    co = doc.add_paragraph(); co.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = co.add_run("Coordinator: prof. univ. dr. Cătălin Boja   ·   Bucharest 2026")
    cr.font.size = Pt(11); cr.font.color.rgb = GREY
    page_break(doc)


# ===========================================================================
# DOCUMENT 1 — Extended Technical & Commercial Report
# ===========================================================================
def build_report():
    doc = Document()
    base_style(doc)
    add_footer_pagenum(doc)

    cover(doc,
          "Architecture and Evaluation of a SIEM Integrating Unsupervised Anomaly "
          "Detection and Knowledge-Grounded Language Models",
          "Extended Technical & Commercial Report",
          "Anomaly ensemble · Hybrid RAG · Agentic CVE ingestion — fully local, fully auditable")

    h(doc, "Table of Contents", 1)
    add_toc(doc)
    page_break(doc)

    # ---- 1. Abstract ----
    h(doc, "1. Abstract", 1)
    p(doc,
      "Traditional Security Information and Event Management (SIEM) platforms rely on predefined, "
      "signature-based rules, which leaves them oblivious to zero-day and polymorphic threats and unable "
      "to convey how anomalous, or how urgent, a given alert really is. This work presents the AI Threat "
      "Engine, an artificial-intelligence layer built on top of the open-source Wazuh SIEM that addresses "
      "these limitations through three complementary components. First, an unsupervised anomaly-detection "
      "ensemble — an Isolation Forest fitted with contamination='auto' and a bottleneck autoencoder "
      "(16→8→4→8→16) trained by iterative self-cleaning — assigns every alert a priority score in [0,100] "
      "from sixteen engineered features, with no human labels required at any stage. Second, a "
      "Retrieval-Augmented Generation (RAG) copilot grounds a locally hosted large language model "
      "(llama3.2 via Ollama) in a curated knowledge base of MITRE ATT&CK techniques, Sigma rules, CISA "
      "vulnerabilities and internal runbooks, retrieved through a hybrid dense (all-MiniLM-L6-v2) and "
      "sparse (BM25) pipeline fused with Reciprocal Rank Fusion, so that every explanation is traceable "
      "to a specific source. Third, a scheduled agentic workflow keeps that knowledge base current by "
      "fetching, de-duplicating, environment-boosting and LLM-scoring new CVEs before routing each "
      "decision into an auditable PostgreSQL ledger. Evaluated on 2,171 real alerts, the ensemble attains "
      "83.4% precision, 95.9% recall, an F1 of 89.2% and an 8.0% false-positive rate, while the RAG layer "
      "reaches an 82.9% retrieval hit-rate and 98.3% citation precision. Crucially, the entire system runs "
      "on-premises with no external API calls, making it suitable for privacy-sensitive and air-gapped "
      "deployments. This report extends the thesis with concrete pathways toward securing mobile and "
      "AI-native applications, engineering for extreme reliability, an adversarial penetration-testing "
      "programme, a guardrail catalogue, a multi-agent evolution, and a go-to-market strategy.")
    page_break(doc)

    # ---- 2. Securing Mobile & AI-System Apps ----
    h(doc, "2. Hardening the Engine to Secure Mobile Apps and AI-System Apps", 1)
    p(doc, "The current engine is host- and network-centric. Two of the fastest-growing attack surfaces — "
           "mobile applications and AI-native (LLM-powered) applications — can be brought under the same "
           "detect-explain-maintain loop with focused extensions. Each subsection is framed as "
           "current state → gap → concrete upgrade.")

    h(doc, "2.1 Securing Mobile Applications", 2)
    p(doc, "Current state. The engine ingests Linux host telemetry (auth.log, syscheck, rootcheck, network "
           "events). It has no visibility into mobile endpoints.")
    p(doc, "Gap. Mobile threats — rooted/jailbroken devices, sideloaded APKs, certificate-pinning bypass, "
           "overlay attacks, malicious SDKs and abused accessibility services — never reach the pipeline.")
    p(doc, "Concrete upgrades:")
    bullet(doc, "ingest MDM/EMM telemetry (Intune, Jamf, MobileIron) plus Android logcat and iOS unified "
                "logs through a Wazuh mobile agent or a syslog forwarder.", "Telemetry ingestion —")
    bullet(doc, "add a mobile feature extractor mirroring the 16-feature design: device-integrity flag "
                "(root/jailbreak), app-signature mismatch, cert-pinning-failure count, permission-escalation "
                "events, install-source (sideload) flag, and off-market package origin.", "Mobile feature set —")
    bullet(doc, "index the MITRE ATT&CK for Mobile matrix and OWASP MASVS/MASTG controls into the RAG "
                "corpus so the copilot maps mobile alerts to techniques such as T1626 (abuse elevation) "
                "and explains the relevant MASVS requirement.", "Knowledge grounding —")
    bullet(doc, "add response playbooks (revoke device, force re-enrolment, block app) as uploaded "
                "runbooks the analyst can be walked through.", "Response —")

    h(doc, "2.2 Securing AI-System (LLM-Powered) Applications", 2)
    p(doc, "Current state. The engine uses an LLM defensively but does not monitor other AI applications.")
    p(doc, "Gap. Organisations increasingly deploy their own LLM apps, which introduce prompt injection, "
           "jailbreaks, training-data exfiltration, model denial-of-wallet and unsafe tool-use — none of "
           "which classical SIEM rules detect.")
    p(doc, "Concrete upgrades:")
    bullet(doc, "treat LLM inference logs (prompts, retrieved context, tool calls, completions) as a new "
                "event source; extract features such as prompt-injection-pattern score, jailbreak-lexicon "
                "hits, output-PII count, tool-call anomaly, and token-burn rate, then score them with the "
                "same ensemble.", "Inference-log monitoring —")
    bullet(doc, "map alerts to the MITRE ATLAS adversarial-ML matrix (e.g., prompt injection, model "
                "evasion, data exfiltration) and index ATLAS into the RAG corpus.", "ATLAS grounding —")
    bullet(doc, "run an output guardrail that scans completions for leaked secrets, unsafe actions and "
                "policy violations before they reach the user, raising a scored alert on violation.",
                "Output scanning —")
    bullet(doc, "detect denial-of-wallet and abuse via anomaly scoring on per-key request-rate and "
                "token-consumption features.", "Abuse detection —")

    # ---- 3. Extreme Reliability ----
    h(doc, "3. Engineering for Extreme Reliability", 1)
    p(doc, "Moving from a working prototype to a system organisations trust in production requires "
           "disciplined reliability engineering across the ML, RAG and data tiers.")
    bullet(doc, "monitor the live score distribution and feature statistics; trigger retraining when the "
                "population shifts (concept drift). The autoencoder's iterative-trimming procedure already "
                "re-estimates the clean distribution from scratch, which makes scheduled retraining safe.",
                "Model-drift detection —")
    bullet(doc, "the ensemble already provides model redundancy — if one detector degrades, the other "
                "still scores. Extend this with a rules-only fallback path when both models are unavailable.",
                "Graceful degradation —")
    bullet(doc, "the CVE agent already fails safe when the LLM is down (routes to the review queue rather "
                "than dropping). Generalise this: wrap every external dependency (Ollama, Qdrant, NVD) in a "
                "circuit breaker with health-checks and exponential backoff.", "Circuit breakers —")
    bullet(doc, "close the ingestion_log run row in a try/finally so a mid-run crash marks the run "
                "'failed' instead of leaving it open — a concrete robustness fix already identified.",
                "Run-state integrity —")
    bullet(doc, "place a durable message queue (Kafka/Redis Streams) between Wazuh logcollector and the "
                "scorer so bursts are buffered and the scorer can scale horizontally behind it.",
                "Horizontal scale —")
    bullet(doc, "run PostgreSQL in a primary/replica high-availability configuration; schedule snapshots "
                "of the Qdrant volume (which is disposable and re-indexable, so recovery is fast).",
                "Data durability —")
    bullet(doc, "add metrics (Prometheus), tracing and structured logs for end-to-end latency, scoring "
                "latency, LLM time-to-first-token and queue depth.", "Observability —")
    p(doc, "Target service levels:")
    table(doc,
          ["Metric", "Target (SLO)", "Rationale"],
          [["Scoring latency (p95)", "< 50 ms/alert", "keeps the live dashboard real-time"],
           ["End-to-end alert→verdict", "< 2 s", "analyst sees scored alerts near-instantly"],
           ["Detection recall", "≥ 95%", "matches the evaluated ensemble"],
           ["False-positive rate", "≤ 8%", "analyst trust threshold from evaluation"],
           ["Availability", "99.9% (scoring path)", "detection must not have blind windows"],
           ["RAG citation precision", "≥ 95%", "grounding must stay trustworthy"]],
          widths=[2.4, 1.9, 2.8])

    page_break(doc)

    # ---- 4. Penetration Testing ----
    h(doc, "4. Penetration-Testing Programme", 1)
    p(doc, "Because the engine is itself a security product, it must be tested adversarially. The programme "
           "below follows a standard methodology (scope → recon → exploitation → reporting) but adds the "
           "AI-specific dimensions that make this system unusual.")

    h(doc, "4.1 Scope and Rules of Engagement", 2)
    p(doc, "Test in an isolated replica of production, never against live customer data. Define in-scope "
           "assets (Flask API, Qdrant, PostgreSQL, Ollama, the CVE fetcher), a testing window, and a "
           "rollback plan. Obtain written authorisation before any test.")

    h(doc, "4.2 Application & Infrastructure Surface", 2)
    bullet(doc, "test every /api/* route for authentication, authorisation and rate-limit bypass; the "
                "current Flask-Limiter and planned bearer-token gateway are the controls under test.",
                "API —")
    bullet(doc, "the CVE fetcher makes outbound requests to NVD/CISA — test for Server-Side Request "
                "Forgery and for poisoned upstream responses.", "SSRF —")
    bullet(doc, "test the document-upload endpoint for path traversal, oversized payloads, and malicious "
                "content that survives into the vector store.", "Upload —")

    h(doc, "4.3 Adversarial Machine-Learning Testing", 2)
    bullet(doc, "attempt low-and-slow attacks that blend into the density (to evade the Isolation Forest) "
                "and manifold-mimicry attacks that reproduce normal feature combinations (to evade the "
                "autoencoder). The thesis argument — that these two evasion strategies are contradictory — "
                "is exactly the hypothesis this test validates.", "Evasion —")
    bullet(doc, "inject crafted benign-looking alerts into the training corpus and measure the shift in "
                "the score threshold; verify iterative trimming discards them.", "Data poisoning —")
    bullet(doc, "attempt membership inference to determine whether a specific alert was in the training "
                "set (a privacy concern for regulated deployments).", "Membership inference —")

    h(doc, "4.4 RAG and LLM Red-Teaming", 2)
    bullet(doc, "upload a document containing injected instructions and verify the system prompt's "
                "grounding constraints hold (the anti-hallucination guards are the control under test).",
                "Indirect prompt injection —")
    bullet(doc, "index adversarial 'intel' and check whether it can be surfaced to mislead an analyst; "
                "test the pending-review quarantine as the mitigation.", "Retrieval poisoning —")
    bullet(doc, "attempt jailbreaks and citation-spoofing (claiming a source that was not retrieved); the "
                "98.3% citation-precision metric is the baseline this test attacks.", "Jailbreak / citation —")
    bullet(doc, "forge or suppress log lines to create false alerts or hide real ones.", "Log injection —")

    h(doc, "4.5 Tooling and Reporting", 2)
    p(doc, "Recommended tooling: Burp Suite and nuclei for the web surface; garak and Microsoft PyRIT for "
           "LLM red-teaming; the Adversarial Robustness Toolbox (ART) for evasion/poisoning; custom scripts "
           "for the feature-space attacks. Report every finding with a severity rating and a remediation.")
    table(doc,
          ["Test area", "Example technique", "Tool", "Severity if unmitigated"],
          [["API", "auth / rate-limit bypass", "Burp, nuclei", "High"],
           ["CVE fetcher", "SSRF, upstream poisoning", "Burp, custom", "High"],
           ["ML evasion", "low-and-slow / manifold mimicry", "ART, custom", "Medium"],
           ["Data poisoning", "training-corpus injection", "custom", "Medium"],
           ["RAG", "indirect prompt injection", "garak, PyRIT", "High"],
           ["LLM", "jailbreak / citation spoof", "garak", "Medium"],
           ["Logs", "alert forging / suppression", "custom", "High"]],
          widths=[1.5, 2.3, 1.4, 1.6])

    page_break(doc)

    # ---- 5. Guardrails ----
    h(doc, "5. Guardrails to Add", 1)
    p(doc, "Guardrails convert the pentest findings into standing controls, organised by pipeline stage.")
    table(doc,
          ["Guardrail", "Threat it blocks", "How to implement"],
          [["Upload sanitisation", "malicious / oversized docs", "type + size limits, content scan, strip active content"],
           ["Source allow-listing", "retrieval poisoning", "only index signed / trusted intel; provenance tags"],
           ["System-prompt hardening", "prompt injection", "answer-only-from-context + explicit refusal instruction"],
           ["Output filtering", "secret / PII leakage", "regex + classifier scan of completions before display"],
           ["Refusal calibration", "hallucinated intel", "unanswerable-query test set; verified 0 over-refusals"],
           ["Structured-output validation", "malformed agent actions", "JSON-schema validation of LLM outputs"],
           ["RBAC + audit logging", "insider misuse", "role-scoped API + immutable action log in Postgres"],
           ["API gateway", "abuse / DoS", "bearer-token auth + Redis rate-limiting (30 rpm)"],
           ["Secrets management", "credential theft", "vault / env isolation; no secrets in code or logs"]],
          widths=[1.9, 1.9, 3.3])

    # ---- 6. More Agentic ----
    h(doc, "6. Making the System More Agentic", 1)
    p(doc, "The present agent performs one bounded LLM decision (CVE relevance scoring) inside a fixed "
           "control flow. That discipline — bounded reasoning, every action logged and reversible — is the "
           "correct foundation on which to add autonomy without losing auditability. The evolution is a "
           "supervised multi-agent SOC:")
    numbered(doc, "reads each scored alert, clusters related alerts, and decides whether escalation is "
                  "warranted — replacing tier-1 queue triage.", "Triage agent —")
    numbered(doc, "for escalated alerts, automatically issues the RAG queries, correlates across the alert "
                  "history, and drafts a hypothesis with citations.", "Investigation agent —")
    numbered(doc, "proposes containment (block IP, isolate host, revoke key) and, within human-approved "
                  "policy, executes it through Wazuh active-response.", "Response agent —")
    numbered(doc, "orchestrates the others, enforces human-in-the-loop gates on any irreversible action, "
                  "and maintains a shared case memory.", "Supervisor —")
    p(doc, "Supporting capabilities: tool-use / function-calling for the response actions; a two-tier "
           "memory (episodic per-case context plus the existing PostgreSQL ledger as durable long-term "
           "memory); an agent-evaluation harness that measures decision accuracy and time-to-containment; "
           "and hard guardrails on autonomy (allow-listed actions, blast-radius limits, mandatory approval "
           "for destructive operations). The governing principle stays constant: agentic does not mean "
           "unbounded.")

    page_break(doc)

    # ---- 7. Go-to-market ----
    h(doc, "7. Go-to-Market: How to Sell It", 1)

    h(doc, "7.1 Who Buys It (Ideal Customer Profiles)", 2)
    table(doc,
          ["Segment", "Why they buy", "Primary value"],
          [["MSSPs / MDR providers", "score + explain across many tenants", "analyst efficiency, margin"],
           ["Mid-market SOCs", "no budget for a big analyst team", "one analyst does tier-1 work"],
           ["Regulated SMBs (GDPR, health, finance)", "data cannot leave premises", "100% local, auditable"],
           ["Defense / government / air-gapped", "no cloud, full sovereignty", "offline, no telemetry"],
           ["AI-product companies", "must secure their own LLM apps", "ATLAS-aware detection"]],
          widths=[2.3, 2.6, 2.2])

    h(doc, "7.2 Value Proposition", 2)
    bullet(doc, "runs entirely on-premises — decisive for privacy-sensitive and air-gapped buyers.", "Local & private —")
    bullet(doc, "local inference means no per-token API bill, unlike cloud copilots.", "No per-token cost —")
    bullet(doc, "every verdict carries a score and every explanation carries a citation.", "Explainable —")
    bullet(doc, "the CVE agent keeps the knowledge base current without manual curation.", "Self-updating —")

    h(doc, "7.3 Pricing Models", 2)
    bullet(doc, "core engine open-source, paid enterprise features (multi-tenant, RBAC, HA).", "Open-core —")
    bullet(doc, "per-monitored-agent or per-node subscription.", "Per-agent —")
    bullet(doc, "support / SLA tiers (community, business, 24×7).", "Support tiers —")
    bullet(doc, "fully managed deployment for buyers without ops capacity.", "Managed —")

    h(doc, "7.4 Step-by-Step Sales Plan", 2)
    numbered(doc, "write a one-sentence positioning statement ('the local, explainable AI copilot for "
                  "Wazuh SOCs') and a crisp problem-solution narrative.", "Position —")
    numbered(doc, "publish a landing page with a two-minute demo video built from the live-demo runbook "
                  "(real attack → detection → grounded explanation).", "Prove —")
    numbered(doc, "recruit 3–5 design partners for free pilots in exchange for feedback and a reference.",
             "Pilot —")
    numbered(doc, "convert one pilot into a written case study with measured outcomes (triage time saved, "
                  "FPs reduced).", "Case study —")
    numbered(doc, "run targeted outbound to the ICP list, leading with the case study.", "Outbound —")
    numbered(doc, "define a pilot-to-paid motion: a 30-day paid pilot with success criteria that convert "
                  "to an annual contract.", "Convert —")
    numbered(doc, "sign MSSP channel partners who resell into their tenant base.", "Channel —")
    numbered(doc, "pursue SOC 2 / ISO 27001 as a sales unlock for regulated buyers.", "Certify —")
    p(doc, "A 90-day launch cadence:")
    table(doc,
          ["Phase", "Weeks", "Milestone"],
          [["Foundation", "1–3", "positioning, landing page, demo video"],
           ["Pilots", "4–8", "3–5 design partners live"],
           ["Proof", "9–10", "first case study published"],
           ["Outbound", "11–12", "ICP outreach + first paid pilot"]],
          widths=[1.8, 1.4, 3.9])

    # ---- 8. Limitations & roadmap ----
    h(doc, "8. Honest Limitations and Roadmap", 1)
    bullet(doc, "the current deployment runs on a single node; enterprise scale needs the queue + HA work "
                "in Section 3.", "Single-node —")
    bullet(doc, "the models were trained on ~2,171 alerts; more data and diversity will improve "
                "generalisation.", "Training-set size —")
    bullet(doc, "single failed PAM logins remain hard (≈45% detection) because they are informationally "
                "identical to a typo; temporal window features are the fix.", "Known edge case —")
    p(doc, "Roadmap priority order: temporal features → multi-agent SOC → mobile & AI-app coverage → "
           "managed offering → certifications.")

    out = os.path.join(HERE, "AI_Threat_Engine_Extended_Report.docx")
    doc.save(out)
    return out


# ===========================================================================
# DOCUMENT 2 — Thesis Improvement & Admissions Guide
# ===========================================================================
def build_admissions():
    doc = Document()
    base_style(doc)
    add_footer_pagenum(doc)

    cover(doc,
          "Elevating the Thesis for Graduate Admissions",
          "ETH Zürich · EPFL · Top European & Australian Programmes",
          "How to strengthen the work, retitle it for academic weight, and position it per destination")

    h(doc, "Table of Contents", 1)
    add_toc(doc)
    page_break(doc)

    h(doc, "1. Strengthening the Thesis Academically", 1)
    p(doc, "Admissions committees at ETH Zürich, EPFL and comparable programmes reward methodological "
           "rigour, reproducibility and honest evaluation over feature count. Each upgrade below is framed "
           "as what committees look for → what the thesis already has → what to add.")

    h(doc, "1.1 Evaluation Rigour", 2)
    p(doc, "Look for: statistically defensible results. Have: precision 83.4%, recall 95.9%, F1 89.2%, a "
           "confusion matrix, a feature ablation and a chunk-size ablation. Add: k-fold cross-validation "
           "with confidence intervals, a precision–recall curve (more honest than ROC on imbalanced data), "
           "significance tests between the baseline and retrained models, and a comparison table against "
           "published SIEM-ML baselines.")

    h(doc, "1.2 Reproducibility", 2)
    p(doc, "Look for: a result a reviewer can rerun. Have: evaluation scripts and datasets in the repo. "
           "Add: fixed random seeds, a dataset card (composition, collection method, licence), pinned "
           "environment (lockfile + container), and a one-command reproduction script.")

    h(doc, "1.3 Threats to Validity", 2)
    p(doc, "Look for: awareness of one's own limitations. Add: an explicit section on internal validity "
           "(the same-family LLM judge biases faithfulness upward — use an independent judge model), "
           "external validity (single deployment, ~2,171 alerts), and construct validity (are the 16 "
           "features measuring what we claim).")

    h(doc, "1.4 Related Work and Formalism", 2)
    p(doc, "Look for: scholarly grounding. Add: a formal problem statement and hypotheses; deeper related "
           "work citing the foundational papers (RAG — Lewis et al. 2020; Isolation Forest — Liu, Ting & "
           "Zhou 2008; sentence embeddings — Reimers & Gurevych 2019) and recent SIEM-ML and LLM-for-"
           "security surveys; and a clear statement of the novel contribution relative to them.")

    h(doc, "1.5 Independent Generation Evaluation", 2)
    p(doc, "The RAG faithfulness score of 0.71 was judged by the same model family that generated the "
           "answers. Re-run the generation evaluation with an independent judge (a different local model, "
           "or human annotation on a sample) and report inter-rater agreement — this single change "
           "materially raises the credibility of the RAG results.")

    page_break(doc)

    h(doc, "2. Better, More Academic Titles", 1)
    p(doc, "Ranked alternatives. The strongest titles foreground the scientific contributions — "
           "self-supervised / unsupervised learning, grounding, explainability, and rigorous evaluation — "
           "rather than the product.")
    table(doc,
          ["#", "Proposed title", "Why it lands"],
          [["1", "Grounded Anomaly Reasoning: Unsupervised Detection and Retrieval-Augmented Explanation "
                 "for Security Operations",
            "pairs the two contributions; 'grounded' + 'reasoning' read as ML-systems research"],
           ["2", "From Signatures to Semantics: A Self-Supervised, Knowledge-Grounded SIEM Architecture",
            "memorable thesis-statement framing; signals the paradigm shift"],
           ["3", "Explainable, Label-Free Threat Detection via Dual-Model Anomaly Ensembles and Local "
                 "Retrieval-Augmented Generation",
            "keyword-dense for search committees; 'label-free' highlights the unsupervised claim"],
           ["4", "Trustworthy SOC Automation: Auditable Anomaly Scoring and Cited Machine Reasoning on "
                 "Open-Source SIEM",
            "'trustworthy' + 'auditable' align with EU trustworthy-AI discourse"],
           ["5", "Detect, Explain, Sustain: A Closed-Loop AI Architecture for Autonomous Security "
                 "Operations",
            "captures the loop; 'autonomous' invites the agentic PhD direction"],
           ["6", "Beyond Rules: Unsupervised Anomaly Ensembles and Grounded Language Models for "
                 "Privacy-Preserving Threat Detection",
            "'privacy-preserving' targets EPFL / EU labs"],
           ["7", "Semantic SIEM: Knowledge-Grounded Language Models for Explainable Intrusion Analysis",
            "concise, coins a term, strong for a paper title"],
           ["8", "Adversarially-Robust, Self-Supervised Threat Detection with Local Knowledge-Grounded "
                 "Reasoning",
            "'adversarially-robust' is the strongest signal for a security-ML committee"]],
          widths=[0.4, 4.0, 3.0])
    p(doc, "Keyword guidance: self-supervised, unsupervised, grounded, explainable, adversarially-robust, "
           "auditable and evaluation all raise the perceived rigour. Avoid product-flavoured words "
           "(engine, copilot, dashboard) in the academic title; keep those for the system name.")

    page_break(doc)

    h(doc, "3. Program-Specific Positioning", 1)
    table(doc,
          ["Destination", "What they value", "Which section to expand", "SOP angle"],
          [["ETH Zürich (systems / security)",
            "rigorous systems evaluation, reproducibility",
            "reliability + adversarial-ML testing",
            "engineering rigour + measurable robustness"],
           ["EPFL (ML security, privacy)",
            "privacy-preserving & self-supervised ML",
            "local/private design + unsupervised training",
            "label-free learning on sensitive data"],
           ["TU Delft / TU Munich / KTH / Imperial",
            "applied security systems with strong eval",
            "the closed detect-explain loop + evaluation",
            "match the specific group's sub-area"],
           ["Australia (UNSW / ANU / CSIRO Data61)",
            "applied, national-security-relevant cyber",
            "agentic SOC + threat-intel automation",
            "operational impact + sovereignty"]],
          widths=[1.9, 1.9, 1.7, 1.9])
    p(doc, "For each application, expand the indicated section, cite one representative paper from the "
           "target group, and mirror the group's vocabulary in the statement of purpose.")

    h(doc, "4. Turning the Thesis into an Application Narrative", 1)
    p(doc, "Structure the statement of purpose and research proposal around a gap → contribution → future "
           "research arc: the gap is that rule-based SIEMs cannot score, explain or self-update; the "
           "contribution is the evaluated, fully-local detect-explain-maintain loop; the future research "
           "is the two hardest, most fundable directions — adversarial robustness of anomaly ensembles, "
           "and safe autonomy in multi-agent SOCs. Propose these as a PhD programme.")
    p(doc, "Publishable sub-papers to extract now:")
    bullet(doc, "the reference-free RAG evaluation methodology (faithfulness + citation precision + "
                "refusal), which is reusable beyond security.", "Methods paper —")
    bullet(doc, "the unsupervised training strategy (contamination='auto' + iterative trimming) as a "
                "label-free anomaly-detection contribution.", "ML paper —")
    bullet(doc, "the governance model for the agentic CVE ingestion (bounded reasoning + auditable "
                "ledger + human-review queue).", "Systems paper —")

    h(doc, "5. Ways It Can Be the Best — Prioritised Checklist", 1)
    p(doc, "Ordered by committee impact per unit of effort:")
    numbered(doc, "add the independent RAG judge and cross-validation with confidence intervals.", "Highest —")
    numbered(doc, "write the threats-to-validity section and a formal problem statement.", "High —")
    numbered(doc, "add the baseline-comparison table and precision–recall curves.", "High —")
    numbered(doc, "publish a reproducibility package (seeds, dataset card, container, one-command rerun).", "Medium —")
    numbered(doc, "retitle using one of the top-3 options and align the abstract to the new title.", "Medium —")
    numbered(doc, "extract one sub-paper and submit to a workshop for an external stamp of quality.", "Stretch —")
    p(doc, "A twelve-week polish plan:")
    table(doc,
          ["Weeks", "Focus", "Output"],
          [["1–3", "evaluation upgrades", "CV + PR curves + independent judge"],
           ["4–6", "scholarly framing", "problem statement, related work, threats to validity"],
           ["7–9", "reproducibility", "public repo + dataset card + container"],
           ["10–12", "positioning", "retitle, rewrite abstract, draft SOP + one sub-paper"]],
          widths=[1.3, 2.4, 3.9])

    out = os.path.join(HERE, "Thesis_Improvement_and_Admissions_Guide.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    r = build_report()
    print("  ✔", os.path.basename(r))
    a = build_admissions()
    print("  ✔", os.path.basename(a))
    print("Done.")
