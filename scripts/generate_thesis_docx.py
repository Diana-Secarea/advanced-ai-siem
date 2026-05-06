#!/usr/bin/env python3
"""
Generează lucrarea de licență în format .docx (~40 pagini, formatare academică).
Rulare:  .venv/bin/python scripts/generate_thesis_docx.py
Ieșire:  docs/Lucrare_AI_Threat_Engine_Wazuh.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Lucrare_AI_Threat_Engine_Wazuh.docx"


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def add_center_title(doc: Document, text: str, size_pt: int = 22, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(12)


def add_subtitle(doc: Document, text: str, size_pt: int = 14, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str, first_indent: bool = False) -> None:
    p = doc.add_paragraph()
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_table_caption(doc: Document, chapter: int, num: int, caption: str) -> None:
    """Tabel X.Y — explicație, stânga, italic (conform modelului)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Tabel {chapter}.{num}. {caption}")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


def add_figure_caption(doc: Document, chapter: int, num: int, caption: str) -> None:
    """Figura X.Y — centrat, italic, sub figură."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figura {chapter}.{num}. {caption}")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)


def add_api_routes_table(doc: Document) -> None:
    """Tabel 5.3 — endpoint-uri principale din backend/server.py."""
    add_table_caption(doc, 5, 3, "Endpoint-uri REST și rolul lor în platformă.")
    rows = [
        ("Cale", "HTTP", "Rol principal"),
        ("/stream", "GET", "Flux SSE de alerte parsate din alerts.log."),
        ("/api/alerts/scored", "GET", "Ultimele alerte JSON cu scor Isolation Forest și etichete."),
        ("/api/alerts/stats", "GET", "Statistici agregate pe niveluri și grupuri."),
        ("/api/chat", "POST", "Chat RAG + alerte locale + apel LLM (bloc)."),
        ("/api/chat/stream", "POST", "Chat cu răspuns LLM în streaming (SSE)."),
        ("/api/chat/sessions", "DELETE", "Ștergere sesiuni chat."),
        ("/api/suggested-benign-rules", "GET", "Listă sugerată de reguli benign."),
        ("/api/benign-rules", "GET/POST", "Citire / adăugare excepții benign."),
        ("/api/benign-rules/<id>", "DELETE", "Eliminare excepție benign."),
        ("/api/suggested-suspicious-groups", "GET", "Grupuri suspecte sugerate."),
        ("/api/suspicious-groups", "GET/POST", "Gestionare grupuri suspecte utilizator."),
        ("/api/suspicious-groups/<name>", "DELETE", "Eliminare grup suspect."),
        ("/api/knowledge/upload", "POST", "Încărcare document în Qdrant (+ Postgres)."),
        ("/api/knowledge/documents", "GET", "Listare documente utilizator din colecție."),
        ("/api/knowledge/documents/<doc_id>", "DELETE", "Ștergere document și vectori."),
        ("/api/vectordb/stats", "GET", "Statistici Qdrant pentru dashboard vectori."),
        ("/, /<path>", "GET", "Servire frontend static."),
    ]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
        if i == 0:
            for c in tbl.rows[i].cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
                _set_cell_shading(c, "E7E6E6")
    doc.add_paragraph()


def add_extra_from_file(doc: Document) -> None:
    """Conținut extins din scripts/thesis_extra_ro.txt (generat sau livrat împreună)."""
    extra_path = Path(__file__).resolve().parent / "thesis_extra_ro.txt"
    if not extra_path.exists():
        add_body(
            doc,
            "Fișierul thesis_extra_ro.txt lipsește; rulați generatorul sau păstrați fișierul lângă script.",
            first_indent=True,
        )
        return
    raw = extra_path.read_text(encoding="utf-8")
    for block in raw.split("\n---\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("SUB|"):
            add_subtitle(doc, block[4:].strip(), size_pt=14)
            continue
        if block.strip() == "__TABLE_API__":
            add_api_routes_table(doc)
            continue
        add_body(doc, block, first_indent=True)


def build_document() -> Document:
    doc = Document()
    set_doc_defaults(doc)

    # --- Pagină titlu ---
    for _ in range(6):
        doc.add_paragraph()
    add_center_title(
        doc,
        "Motor AI pentru analiza amenințărilor în Wazuh:\n"
        "detecție prin Isolation Forest și pipeline RAG hibrid\n"
        "(backend Flask, Qdrant, PostgreSQL)",
        size_pt=16,
    )
    doc.add_paragraph()
    add_center_title(doc, "Lucrare de licență", size_pt=14, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "Autor: [Nume Prenume]",
        "Coordonator: [Titlu Dr. Nume Prenume]",
        "Universitatea: [Denumire]",
        "Facultatea: [Denumire]",
        "Program de studii: [Denumire]",
        "An universitar: 2025–2026",
    ):
        r = p.add_run(line + "\n")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    page_break(doc)

    # --- Cuprins (actualizare manuală sau din Word: Referințe → Cuprins) ---
    add_center_title(doc, "Cuprins", size_pt=14)
    cuprins = """1. Introducere
2. Inteligența artificială în securitatea cibernetică: domeniul problemei
   2.1. Peisajul tot mai complex al amenințărilor
   2.2. Avantajele AI în cybersecurity
   2.3. Direcții viitoare în securitatea asistată de AI
   2.4. SIEM și rolul său operațional
   2.5. Wazuh ca platformă SIEM
   2.6. Limitările SIEM tradițional și ale Wazuh
3. Metode și algoritmi utilizați
   3.1. Isolation Forest pentru detecția anomaliilor
   3.2. RAG (Retrieval-Augmented Generation)
   3.3. Căutare hibridă: semantică + BM25
   3.4. Embedding-uri și indexare vectorială
   3.5. Re-ranking (cross-encoder) — perspective
4. Arhitectura și proiectarea soluției
5. Implementarea soluției
6. Concluzii
Bibliografie
Anexe"""
    for line in cuprins.split("\n"):
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.3
    note = (
        "Notă de formatare: după inserarea titlurilor finale în Word, utilizați "
        "Referințe → Cuprins pentru generarea automată a cuprinsului. "
        "Numerotarea paginilor poate începe la capitolul 1 (Introducere), "
        "cu prima pagină și cuprinsul fără număr, conform îndrumarului facultății."
    )
    add_body(doc, note)
    page_break(doc)

    # ===================== CAP 1 =====================
    add_center_title(doc, "1. Introducere")
    intro = [
        "Organizațiile moderne operează într-un mediu în care volumul de telemetrie de securitate "
        "crește continuu, iar actorii malitioși folosesc tehnici din ce în ce mai automate și "
        "mai greu de semnalizat prin reguli statice. Platformele de tip SIEM (Security Information "
        "and Event Management) centralizează evenimentele, aplică reguli și generează alerte, "
        "însă echipele SOC (Security Operations Center) rămân confruntate cu două probleme structurale: "
        "supraîncărcarea cu alerte și lipsa contextualizării rapide a unui eveniment izolat în raport "
        "cu cunoașterea despre amenințări (tactici MITRE, reguli Sigma, IOC-uri, CVE-uri exploatabile etc.).",

        "Această lucrare prezintă proiectarea și implementarea unui sistem extins de analiză, "
        "denumit în continuare AI Threat Engine, integrat cu Wazuh și cu un backend web (Flask) "
        "care expune capabilități de monitorizare, chat asistat de RAG și gestionare a excepțiilor "
        "operaționale (reguli considerate benign de către analist). Soluția combină două paradigme "
        "complementare de inteligență artificială aplicată în securitate: (1) detecția anomaliilor "
        "printr-un model nesupervizat de tip Isolation Forest, antrenat pe alerte considerate "
        "operational „curate”, și (2) un pipeline RAG (Retrieval-Augmented Generation) care "
        "recuperează fragmente relevante dintr-o bază de cunoștințe de threat intelligence și "
        "le furnizează unui model de limbaj (LLM) pentru sinteză și recomandări.",

        "Motivația alegerii temei derivă din necesitatea practică de a reduce timpul de triaj și "
        "de a îmbunătăți calitatea deciziilor fără a înlocui expertiza umană. Un scor comportamental "
        "numeric (0–100) permite prioritizarea alertelor, iar RAG limitează halucinațiile specifice "
        "LLM-urilor prin ancorarea răspunsurilor în documente recuperate și în alerte reale din "
        "sesiunea de monitorizare.",

        "Obiectivele lucrării sunt: (O1) definirea clară a problemei de domeniu și a limitelor "
        "SIEM-ului clasic; (O2) descrierea riguroasă a algoritmilor (Isolation Forest, RAG, fuziune "
        "RRF pentru recuperare hibridă); (O3) proiectarea arhitecturii pe straturi (telemetrie, ML, "
        "cunoaștere, aplicație, persistență); (O4) implementarea și integrarea cu Wazuh și backend-ul; "
        "(O5) evaluarea modelului de anomalii în două faze — model de referință (baseline) și model "
        "reantrenat — pe seturi de antrenare și test.",

        "Contribuțiile practice ale implementării includ: extragerea unui set bogat de caracteristici "
        "numerice din alertele Wazuh (inclusiv semnale agnostice față de ID-ul regulii, pentru a prinde "
        "atacuri necunoscute), calibrarea scorului brut al Isolation Forest într-o scară interpretabilă, "
        "stabilirea pragului de anomalie din distribuția alertelor benign pe setul de antrenare, "
        "integrarea unei arhitecturi de producție cu Qdrant (căutare densă + sparse) și PostgreSQL "
        "(metadate), precum și mecanisme de feedback uman (liste de reguli benign și grupuri suspecte).",

        "Avantajele soluției propuse rezidă în capacitatea de a detecta comportamente atipice fără "
        "etichete complete pentru toate tipurile de atac, în paralel cu furnizarea de context "
        "inteligibil analistului. Dezavantajele includ dependența de calitatea datelor de bază benign, "
        "riscul de drift operațional (schimbări în infrastructură), și necesitatea reantrenării periodice. "
        "Limitările sunt discutate explicit în capitolul final.",

        "Structura lucrării urmează logica prezentării progresive a soluției: după introducere, "
        "capitolul 2 analizează domeniul și poziționează Wazuh în ecosistemul SIEM; capitolul 3 "
        "prezintă fundamentul teoretic și algoritmic; capitolul 4 descrie arhitectura; capitolul 5 "
        "detaliază implementarea și rezultatele experimentale; capitolul 6 sintetizează concluziile "
        "și direcțiile de cercetare ulterioară.",
    ]
    for para in intro:
        add_body(doc, para, first_indent=True)

    page_break(doc)

    # ===================== CAP 2 =====================
    add_center_title(doc, "2. Inteligența artificială în securitatea cibernetică: domeniul problemei")

    add_subtitle(doc, "2.1. Peisajul tot mai complex al amenințărilor")
    c21 = [
        "Atacurile cibernetice moderne imbrică adesea mai multe faze: acces inițial, persistență, "
        "mișcare laterală, exfiltrare sau impact. Fiecare fază poate genera evenimente care, "
        "izolate, par banale (ex.: o autentificare eșuată, o cerere HTTP neobișnuită). Din perspectiva "
        "SOC, dificultatea nu este doar detectarea unui indicator cunoscut, ci construirea unei "
        "povești coerente din fragmente eterogene de telemetrie.",

        "În același timp, automatizarea atacurilor (scanări în masă, credential stuffing, exploatări "
        "de vulnerabilități cunoscute) generează volume mari de alerte cu grad variabil de severitate. "
        "Fără mecanisme de prioritizare, analiștii petrec timp disproporționat pe evenimente de nivel "
        "scăzut, ceea ce crește latența de răspuns pentru incidentele critice.",

        "Literatura și rapoartele industriale subliniază creșterea sofisticării și a impactului "
        "economic al incidentelor, precum și importanța combinației între prevenție, detecție și "
        "răspuns orchestrat. În acest context, instrumentele care reduc zgomotul și adaugă context "
        "devin parte a infrastructurii de securitate, nu simple „nice-to-have”.",
    ]
    for para in c21:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "2.2. Avantajele AI în cybersecurity")
    c22 = [
        "Aplicarea AI în securitate nu înseamnă înlocuirea regulilor, ci completarea lor. Modelele "
        "nesupervizate pot evidenția tipare statistice rare sau combinații neobișnuite de atribute, "
        "în timp ce modelele supervizate pot clasifica atacuri cunoscute dacă există date etichetate "
        "suficiente. În practică, etichetarea completă este costisitoare; de aceea, abordarea "
        "„antrenare pe benign + evaluare pe atac” este frecvent utilizată în prototipuri operaționale.",

        "RAG adaugă un al doilea strat: recuperarea de documente relevante permite modelului de limbaj "
        "să rămână ancorat în surse explicite (reguli Sigma, tehnici ATT&CK, buletine CISA KEV etc.). "
        "Aceasta îmbunătățește trasabilitatea răspunsurilor și reduce tendința de a inventa conexiuni "
        "între evenimente fără suport documentar.",

        "Un avantaj esențial al integrării AI în SIEM este scalabilitatea cognitivă: un asistent "
        "poate sintetiza rapid zeci de surse, în timp ce analistul validează ipotezele și decide "
        "acțiunea. Totuși, succesul depinde de guvernanță: politici de confidențialitate, controlul "
        "accesului la date și auditarea deciziilor automate.",
    ]
    for para in c22:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "2.3. Direcții viitoare în securitatea asistată de AI")
    c23 = [
        "Direcțiile de evoluție includ: învățare continuă sau periodică cu feedback uman, explicabilitate "
        "îmbunătățită a scorurilor, corelare pe graf între entități (utilizatori, hosturi, adrese IP), "
        "și integrarea semnalelor din mai multe surse (EDR, IDS, cloud audit logs). Pentru mediile "
        "enterprise, se urmărește reducerea costului de mentenanță prin pipeline-uri MLOps și "
        "monitorizarea deriva modelului.",

        "În cadrul proiectului de față, unele dintre aceste direcții sunt pregătite conceptual: "
        "există scripturi de colectare a datelor zilnice, mecanisme de reantrenare, endpoint-uri "
        "pentru excepții utilizator și un dashboard pentru vector store. Extensiile ulterioare pot "
        "include re-ranking cu cross-encoder și învățare semi-supervizată pe baza deciziilor analistului.",
    ]
    for para in c23:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "2.4. SIEM și rolul său operațional")
    c24 = [
        "Un SIEM agregă loguri, le normalizează, aplică reguli de corelare și produce alerte. "
        "Funcțiile tipice includ: căutare și investigație, raportare pentru conformitate, integrare cu "
        "sisteme de ticketing și orchestrare a răspunsului. În literatura de specialitate, SIEM-ul "
        "este adesea descris ca o componentă centrală a operațiunilor de securitate, dar și ca o "
        "sursă de volum mare de evenimente care trebuie filtrate inteligent.",

        "Limitarea fundamentală a SIEM-ului bazat exclusiv pe reguli este că regulile exprimă ceea ce "
        "deja cunoaștem. Atacatorii pot ocoli semnături cunoscute sau pot acționa sub pragurile care "
        "declanșează corelări. De aceea, un strat probabilistic (scor) și un strat semantic (RAG) "
        "pot îmbunătăți acoperirea fără a elimina auditabilitatea regulilor deterministe.",
    ]
    for para in c24:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "2.5. Wazuh ca platformă SIEM")
    add_subtitle(doc, "2.5.1. Motorul de detecție și ierarhia regulilor", size_pt=12)
    c251 = [
        "Wazuh produce alerte structurate în JSON, incluzând identificatorul și nivelul regulii, "
        "descrierea, grupurile, etichetele MITRE (când sunt disponibile), agentul sursă și câmpurile "
        "din payload (de exemplu srcip, srcport, utilizator, comenzi). Această structură este "
        "esențială pentru extragerea automatizată de caracteristici și pentru interogări ulterioare.",

        "Totuși, prezența etichetelor MITRE nu echivalează întotdeauna cu un incident malign: multe "
        "evenimente de rutină pot fi cartografiate la tehnici pentru claritate tactică. Din acest "
        "motiv, în implementarea de față, etichetele MITRE nu sunt folosite direct ca semnal de "
        "etichetă „atac” pentru antrenare, pentru a evita contaminarea setului benign.",
    ]
    for para in c251:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "2.5.2. Modulele Wazuh relevante proiectului", size_pt=12)
    add_body(
        doc,
        "Proiectul consumă alertele agregate (fișiere tip alerts.json / alerts.log) și le procesează "
        "în backend pentru: (a) afișare în interfață, (b) scorizare cu Isolation Forest, (c) căutare "
        "lexicală în sesiunea curentă de monitorizare, (d) augmentarea mesajului către LLM cu fragmente "
        "din threat intelligence recuperate din Qdrant.",
        first_indent=True,
    )

    add_subtitle(doc, "2.5.3. Structura unei alerte reale", size_pt=12)
    add_body(
        doc,
        "O alertă tipică conține timestamp, rule.id, rule.level, rule.description, rule.groups, "
        "câmpul full_log (text brut foarte util pentru cuvinte-cheie precum failed/denied), secțiunea "
        "data cu câmpuri specifice sursei, și metadata despre agent. Această heterogenitate impune "
        "o inginerie atentă a caracteristicilor: unele semnale trebuie citite din full_log, nu doar "
        "din mesajul scurt, pentru a reflecta realitatea operațională.",
        first_indent=True,
    )

    add_subtitle(doc, "2.5.4. Analiza AI Threat Engine asupra alertei", size_pt=12)
    add_body(
        doc,
        "Pipeline-ul aplicat unei alerte include: extragerea vectorului de caracteristici, scalarea, "
        "aplicarea decision_function din Isolation Forest, normalizarea în 0–100, compararea cu "
        "pragul învățat, aplicarea excepțiilor utilizator (reguli benign), etichetarea pentru "
        "afișare (NORMAL / SUSPICIOUS / HIGH / BENIGN în API-ul de listare), și — în modul chat — "
        "interogarea bazei de cunoștințe și construirea contextului pentru LLM.",
        first_indent=True,
    )

    add_subtitle(doc, "2.6. Limitările SIEM tradițional și ale Wazuh")
    c26 = [
        "Printre limitări se numără: dependența de tuning-ul regulilor, riscul de false pozitive "
        "la evenimente administrative legitime cu nivel ridicat, dificultatea generalizării la atacuri "
        "noi și costul investigării manuale. Wazuh, ca platformă open-source, aduce avantaje de "
        "transparență și extensibilitate, dar necesită integrare suplimentară pentru capabilități "
        "de tip „AI product” (vector DB, orchestrare LLM, managementul modelelor).",

        "Soluția propusă tratează explicit unele capcane practice: excluderea unor intervale CDN "
        "pentru a nu marca traficul legitim prin proxy ca „IP extern suspect”, plafonarea unor "
        "caracteristici (ex. număr de cuvinte) pentru a nu confunde rapoarte verbose benign cu "
        "loguri de atac complexe, și separarea seturilor train/test pentru a măsura generalizarea.",
    ]
    for para in c26:
        add_body(doc, para, first_indent=True)

    page_break(doc)

    # ===================== CAP 3 =====================
    add_center_title(doc, "3. Metode și algoritmi utilizați")

    add_subtitle(doc, "3.1. Isolation Forest pentru detecția anomaliilor")
    add_subtitle(doc, "3.1.1. Rolul său în detectarea amenințărilor necunoscute", size_pt=12)
    c311 = [
        "Isolation Forest construiește o pădure de arbori de izolare prin partitionări aleatoare ale "
        "spațiului caracteristicilor. Intuiția fundamentală este că punctele anomale sunt mai ușor "
        "de izolat (necesită mai puține split-uri) decât punctele normale, aflate dense în interiorul "
        "distribuției. Algoritmul este atractiv în SOC deoarece nu presupune o formă parametrică "
        "strictă a distribuției datelor și funcționează bine în dimensiuni moderate, cu precauții "
        "de preprocesare.",

        "În această lucrare, modelul este antrenat pe alerte clasificate ca benign în etapa de "
        "pregătire a datelor. Alertelor de tip atac li se calculează scorul post-antrenare, așteptându-se "
        "ca acestea să iasă în evidență față de baza de referință. Această strategie este aliniată "
        "cu practica „one-class learning” în care clasa normală este bine reprezentată, iar anomaliile "
        "sunt rare și eterogene.",
    ]
    for para in c311:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "3.1.2. Comparație cu metode alternative", size_pt=12)
    add_body(
        doc,
        "Alternativele includ One-Class SVM, Local Outlier Factor, autoencoderi și clasificatori "
        "supervizate (Random Forest, gradient boosting). Supervizarea completă necesită etichete "
        "fiabile pentru toate clasele; în datele Wazuh, etichetarea automată prin reguli și cuvinte "
        "cheie este utilă, dar imperfectă — de aceea evaluarea include atât metrici pe setul de "
        "antrenare, cât și pe un set de test separat, pentru a observa stabilitatea și supra-adaptarea.",
        first_indent=True,
    )

    add_subtitle(doc, "3.1.3. Colectarea și strategia datelor de antrenare", size_pt=12)
    c313 = [
        "Strategia aplicată în codul proiectului include: agregarea alertelor în fișiere combinate, "
        "colectarea zilnică opțională, separarea alertelor „curate” de cele folosite la evaluarea "
        "atacurilor și păstrarea unui fișier de test distinct (all_test_alerts.json) pentru "
        "măsurători după reantrenare.",

        "Excepțiile definite de utilizator în backend (benign_rules.json) influențează atât "
        "antrenarea, cât și inferența: o regulă marcată benign nu mai este tratată ca atac în "
        "construirea ground truth-ului pentru rapoarte, iar scorul poate fi plafonat pentru a reduce "
        "alarmele false repetitive.",
    ]
    for para in c313:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "3.2. RAG (Retrieval-Augmented Generation)")
    c32 = [
        "RAG combină două module: un modul de recuperare (retrieval) care returnează documente "
        "relevante interogării, și un modul generator (LLM) care produce răspunsul condiționat de "
        "aceste documente. În aplicația de față, recuperarea se face din colecția Qdrant aferentă "
        "threat intelligence-ului, iar generatorul este un LLM servit local prin API compatibil "
        "Ollama.",

        "Un principiu important implementat în backend este „porțile” de context: dacă toate alertele "
        "potrivite interogării au nivel scăzut, threat intelligence-ul nu este injectat în prompt, "
        "pentru a evita asocierea forțată a evenimentelor de rutină cu tactici de atac.",
    ]
    for para in c32:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "3.3. Căutare hibridă: semantică + BM25 (sparse)")
    c33 = [
        "Căutarea exclusiv semantică poate rata potrivirile lexicale exacte (nume CVE, string-uri "
        "unice, identificatori). Recuperarea exclusiv lexicală poate rata sinonimia și reformulările. "
        "De aceea, sistemul folosește două reprezentări: un vector dens (similaritate cosinus) și un "
        "vector sparse asociat unui model BM25-like prin fastembed, urmate de fuziunea listelor de "
        "candidați prin Reciprocal Rank Fusion (RRF).",

        "RRF acumulează contribuții inverse proporționale cu rangul în fiecare listă, astfel încât "
        "documentele care apar bine în ambele moduri de căutare urcă în clasament. Constanta de "
        "netezire k (implicit 60) controlează influența rangurilor foarte mari față de cele mici.",
    ]
    for para in c33:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "3.4. Embedding-uri și indexare vectorială")
    c34 = [
        "Embedding-urile dense transformă textul episoadelor (rezumat + etichete + entități) într-un "
        "spațiu vectorial unde proximitatea semnifică similaritate semantică. Stocarea și căutarea "
        "acestor vectori în Qdrant permite scalarea către sute de mii/milioane de puncte, cu "
        "filtrare pe metadate (source_type, tags, episode_type).",

        "PostgreSQL stochează în paralel înregistrări structurate (JSONB) utile pentru audit, "
        "interogări SQL și reîncărcare controlată a datelor, în timp ce Qdrant optimizează latența "
        "de interogare pentru RAG.",
    ]
    for para in c34:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "3.5. Re-ranking (cross-encoder) — perspective")
    c35 = [
        "Re-ranking-ul cu cross-encoder evaluează perechi (interogare, document) cu un model "
        "bidirecțional, oferind adesea o precizie mai mare la top-k decât similaritatea cosinus "
        "dintre embedding-uri create separat. În versiunea actuală a codului, precizia la top-k "
        "este îmbunătățită prin prefetch suficient de mare și RRF; adăugarea unui cross-encoder "
        "este o extensie naturală pentru medii cu cerințe stricte de calitate a recuperării.",

        "Tabelul 3.1 sintetizează pașii pipeline-ului de căutare hibridă.",
    ]
    for para in c35:
        add_body(doc, para, first_indent=True)

    add_table_caption(doc, 3, 1, "Pașii pipeline-ului de recuperare hibridă (dens + sparse + RRF).")
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    hdr = ("Pas", "Descriere")
    for j, h in enumerate(hdr):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(cell, "E7E6E6")
    rows = [
        ("P1", "Construirea textului episodului din summary, tags și entități."),
        ("P2", "Codare densă cu SentenceTransformer și normalizare."),
        ("P3", "Codare sparse (BM25-like) pentru potriviri lexicale."),
        ("P4", "Prefetch candidați (ex. 40) din fiecare canal + filtre metadata."),
        ("P5", "Fuziune RRF și returnare top-k rezultate către stratul aplicație."),
    ]
    for i, (a, b) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    doc.add_paragraph()

    page_break(doc)

    # ===================== CAP 4 =====================
    add_center_title(doc, "4. Arhitectura și proiectarea soluției")

    add_subtitle(doc, "4.1. Prezentare generală a arhitecturii")
    c41 = [
        "Arhitectura este organizată pe straturi: (1) sursa de telemetrie Wazuh; (2) motorul de "
        "caracteristici și Isolation Forest în pachetul ai_engine; (3) subsistemul RAG (ingestion, "
        "indexare, QdrantStore); (4) backend Flask care orchestrează API-urile, limitarea ratelor "
        "și integrarea LLM; (5) frontend static pentru dashboard și chat; (6) containere Docker "
        "pentru Qdrant și PostgreSQL.",

        "Separarea responsabilităților este importantă: RAGCoreSystem gestionează operațiuni de "
        "ingest/index/search, în timp ce scorizarea și pragurile de anomalie rămân în AnomalyDetector. "
        "Aceasta permite înlocuirea sau upgrade-ul unui strat fără a invalida întregul sistem.",
    ]
    for para in c41:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "4.2. Diagrama componentelor (descriere pentru Figura 4.1)")
    add_figure_caption(
        doc,
        4,
        1,
        "Diagrama componentelor: Wazuh → backend (scor + căutare alerte) → Qdrant/PostgreSQL → LLM. "
        "Inserați în Word diagrama realizată în draw.io / Lucidchart conform acestei descrieri.",
    )
    add_body(
        doc,
        "Componente principale și fluxuri: (C1) Agent Wazuh și managerul emit alerte JSONL; (C2) "
        "backend-ul citește incremental alerts.json pentru sesiune; (C3) AnomalyDetector încarcă "
        "modelul .pkl (cale producție sau locală); (C4) QdrantStore realizează interogări hibrid; "
        "(C5) endpoint-ul /api/chat combină context TI + alerte și apelează LLM; (C6) utilizatorul "
        "poate încărca documente noi în bază prin /api/knowledge/upload.",
        first_indent=True,
    )

    add_subtitle(doc, "4.3. Fluxul datelor și pipeline-ul de procesare")
    add_figure_caption(
        doc,
        4,
        2,
        "Fluxul pentru /api/chat: interogare → căutare alerte → scor IF → (opțional) căutare Qdrant → "
        "construire prompt → răspuns LLM. Inserați fluxul sub formă de organigramă în lucrarea finală.",
    )
    c43 = [
        "Fluxul de chat implementat include pași determinați de cod: căutarea în memorie a alertelor "
        "recente prin potrivire lexicală, scorizarea alertelor găsite, decizia de a include sau nu "
        "threat intelligence-ul în funcție de nivelul maxim al alertelor potrivite, recuperarea "
        "fragmentelor din Qdrant, concatenarea contextului și apelul LLM (streaming sau bloc).",

        "Fluxul pentru monitorizare include SSE din alerts.log (parsare blocuri) și API-uri JSON "
        "pentru statistici și alerte scorificate.",
    ]
    for para in c43:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "4.4. Structura bazei de cunoștințe RAG")
    c44 = [
        "Baza de cunoștințe include episoade generate din surse multiple: tehnici MITRE ATT&CK, "
        "reguli YARA, reguli Sigma, catalog CISA KEV, controale CIS, D3FEND, grupuri și software "
        "MITRE, pulsuri OTX (dacă este configurată cheia API), advisory-uri vendor și încărcări "
        "utilizator. Fiecare înregistrare devine un payload Qdrant cu câmpuri indexate pentru filtrare.",

        "Episoadele de alertă (colecție separată) permit căutarea de similitudine între incidente, "
        "utilă pentru triaj comparativ și pentru construirea cazurilor repetitive.",
    ]
    for para in c44:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "4.5. Proiectarea bazei de date și stocare")
    c45 = [
        "PostgreSQL păstrează înregistrări normalizate (id, source_type, name, description, tags, "
        "raw_data JSONB) și jurnalizează ingestia. Qdrant păstrează vectorii și metadatele pentru "
        "căutare rapidă. Această separare urmează principiul poliglotei de persistență: fiecărui "
        "tip de interogare îi corespunde un motor optimizat.",

        "Docker Compose fixează versiunile serviciilor, volumele persistente și healthcheck-urile, "
        "ceea ce sprijină reproductibilitatea experimentelor și a demonstrațiilor.",
    ]
    for para in c45:
        add_body(doc, para, first_indent=True)

    page_break(doc)

    # ===================== CAP 5 =====================
    add_center_title(doc, "5. Implementarea soluției")

    add_subtitle(doc, "5.1. Stiva tehnologică")
    c51 = [
        "Implementarea principală este în Python. Backend-ul Flask expune rute pentru streaming, "
        "chat, scoruri, încărcare documente și administrarea listelor benign/suspicious. Motorul "
        "ML folosește scikit-learn (Isolation Forest, StandardScaler) și serializare joblib. "
        "Subsistemul RAG folosește sentence-transformers, fastembed, qdrant-client și driverul "
        "PostgreSQL.",
    ]
    for para in c51:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "5.2. Pipeline-ul de antrenare Isolation Forest")
    c52 = [
        "Scriptul de antrenare încarcă alerte JSONL, le separă în benign vs atac folosind aceeași "
        "logică ca la evaluare (attack_labels.py), antrenează pe benign, calibrează scor_min/scor_max "
        "pe percentile, calculează pragul ca percentile a scorurilor normalizate pe benign (cu "
        "minimum 40), salvează modelul și evaluează pe atacuri și pe benign pentru rata de detecție "
        "și false pozitive.",

        "Caracteristicile extrase includ atât semnale structurale (nivel, ID regulă, număr câmpuri "
        "în data), cât și semnale comportamentale: IP sursă extern (cu excluderi CDN), pattern-uri "
        "URL suspecte, semnale de utilizator necunoscut și indicatori pentru modificări privilegiate "
        "care altfel ar arăta ca evenimente de rutină.",
    ]
    for para in c52:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "5.3. Implementarea sistemului RAG")
    c53 = [
        "Ingestia populează directorul threat_intel/ prin module dedicate. Indexerul qdrant_indexer "
        "citește fișierele JSON, construiește episoade, face upsert în PostgreSQL și indexare în "
        "Qdrant pe sursă, evitând încărcarea întregului corpus în RAM odată.",

        "La interogare, QdrantStore encodă interogarea, prefetch-ează candidați din ambele spații "
        "vectoriale și fuzionează prin RRF. Rezultatele sunt mapate într-un format comun pentru "
        "Pattern Analyzer / afișare.",
    ]
    for para in c53:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "5.4. Integrarea cu Wazuh")
    add_body(
        doc,
        "Integrarea presupune acces la fișierele de alertă ale managerului Wazuh. În medii reale, "
        "backend-ul poate rula cu drepturi adecvate pentru citirea /var/ossec/logs/alerts/. În "
        "mediul de dezvoltare se pot folosi fișiere locale. Modelul antrenat poate fi scris în "
        "/var/ossec/ai_models/ sau în copie locală pentru teste.",
        first_indent=True,
    )

    add_subtitle(doc, "5.5. Monitorizarea alertelor și procesare în timp real")
    add_body(
        doc,
        "Sesiunea de monitorizare pentru alerts.json pornește de la sfârșitul fișierului pentru a "
        "evita încărcarea istoricului lung în memorie; doar liniile noi apendate sunt colectate. "
        "Aceasta este o decizie de proiect pentru dashboard „live”, cu cost redus de I/O și memorie.",
        first_indent=True,
    )

    add_subtitle(doc, "5.6. Simulare de atac și testare")
    add_body(
        doc,
        "Testarea include colectarea unui set de test după antrenare, scripturi de simulare și "
        "rapoarte de evaluare (confuzie, distribuții de scor, interpretare per regulă). Se "
        "recomandă păstrarea fișierelor de ieșire și a figurilor generate pentru încărcare în anexă.",
        first_indent=True,
    )

    add_subtitle(doc, "5.7. Mecanismul de scorizare Isolation Forest (detaliu tehnic)")
    c57 = [
        "Fie x vectorul de caracteristici al unei alerte. După scalare z-score pe fiecare dimensiune, "
        "modelul returnează s = decision_function(x). Valorile mai mici ale lui s indică anomalie "
        "mai pronunțată. Pentru interpretabilitate, se calculează limitele [s_min, s_max] pe baza "
        "percentilelor scorurilor brute din datele de antrenare benign, apoi se mapează invers la "
        "scor normalizat N ∈ [0,100], astfel încât N mare înseamnă comportament mai anormal.",

        "Pragul T este derivat din distribuția scorurilor normalizate pe benign (ex. percentile 90), "
        "cu un prag minim pentru a evita sensibilitatea excesivă. Decizia finală este N ≥ T. "
        "Excepțiile utilizatorului pentru reguli benign pot plafona N pentru a reflecta cunoștința "
        "domeniului care nu este încă învățată suficient din date.",

        "Scorul combinat multi-dimensional (timp, frecvență, rețea) folosit în unele rapoarte "
        "ponderizează scorul de anomalie cu semnale euristice suplimentare pentru prioritizare "
        "operațională; acesta este complementar, nu înlocuitor al pragului IF.",
    ]
    for para in c57:
        add_body(doc, para, first_indent=True)

    add_subtitle(doc, "5.8. Evaluare inițială (model baseline) și reantrenare — rezultate")
    add_body(
        doc,
        "Experimentul raportat măsoară performanța pe două seturi: (D_train) date combinate de "
        "antrenare cu 1269 alerte benign și 603 alerte etichetate atac; (D_test) set de test cu "
        "171 alerte benign și 11 alerte atac. Etichetarea urmărește aceeași funcție is_attack_alert "
        "ca în codul proiectului, cu excepțiile utilizatorului de tip benign dezactivate în "
        "experimentul reproducibil local (set gol), pentru consistență cu raportul generat automat.",
        first_indent=True,
    )

    add_table_caption(doc, 5, 1, "Rezultate evaluare Isolation Forest — set antrenare (D_train).")
    t1 = doc.add_table(rows=3, cols=8)
    t1.style = "Table Grid"
    h1 = ["Model", "Sep. scor", "TPR%", "FPR%", "Prec%", "Rec%", "F1%", "Prag"]
    for j, h in enumerate(h1):
        c = t1.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E6E6")
    t1.rows[1].cells[0].text = "Baseline"
    t1.rows[1].cells[1].text = "75.58"
    t1.rows[1].cells[2].text = "95.69"
    t1.rows[1].cells[3].text = "8.83"
    t1.rows[1].cells[4].text = "83.74"
    t1.rows[1].cells[5].text = "95.69"
    t1.rows[1].cells[6].text = "89.32"
    t1.rows[1].cells[7].text = "49"
    t1.rows[2].cells[0].text = "Reantrenat"
    t1.rows[2].cells[1].text = "76.15"
    t1.rows[2].cells[2].text = "95.19"
    t1.rows[2].cells[3].text = "8.20"
    t1.rows[2].cells[4].text = "84.66"
    t1.rows[2].cells[5].text = "95.19"
    t1.rows[2].cells[6].text = "89.62"
    t1.rows[2].cells[7].text = "51"
    doc.add_paragraph()

    add_table_caption(doc, 5, 2, "Rezultate evaluare Isolation Forest — set test (D_test).")
    t2 = doc.add_table(rows=3, cols=8)
    t2.style = "Table Grid"
    for j, h in enumerate(h1):
        c = t2.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(c, "E7E6E6")
    t2.rows[1].cells[0].text = "Baseline"
    t2.rows[1].cells[1].text = "56.13"
    t2.rows[1].cells[2].text = "81.82"
    t2.rows[1].cells[3].text = "4.09"
    t2.rows[1].cells[4].text = "56.25"
    t2.rows[1].cells[5].text = "81.82"
    t2.rows[1].cells[6].text = "66.67"
    t2.rows[1].cells[7].text = "49"
    t2.rows[2].cells[0].text = "Reantrenat"
    t2.rows[2].cells[1].text = "59.85"
    t2.rows[2].cells[2].text = "81.82"
    t2.rows[2].cells[3].text = "4.09"
    t2.rows[2].cells[4].text = "56.25"
    t2.rows[2].cells[5].text = "81.82"
    t2.rows[2].cells[6].text = "66.67"
    t2.rows[2].cells[7].text = "51"
    doc.add_paragraph()

    add_subtitle(doc, "5.9. Interpretarea comparativă")
    c59 = [
        "Pe setul D_train, reantrenarea îmbunătățește separarea medie a scorurilor și reduce rata "
        "falselor pozitive, cu o ușoară scădere a TPR care rămâne foarte ridicată (95.19%). Pragul "
        "crește de la 49 la 51, reflectând o calibrare mai conservatoare față de distribuția benign.",

        "Pe setul D_test, TPR rămâne identic (81.82%) pentru ambele modele, iar FPR rămâne 4.09%, "
        "ceea ce indică stabilitatea generalizării pe eșantionul mic de atacuri (11). Precizia "
        "moderată pe test este consecința dezechilibrului claselor și a numărului mic de exemple "
        "pozitive: metricile trebuie interpretate cu prudență statistică.",

        "Concluzia practică este că reantrenarea aduce îmbunătățiri măsurabile pe datele mari de "
        "antrenare fără a degradează performanța pe test în acest experiment, ceea ce validează "
        "strategia de antrenare pe benign și calibrare periodică.",
    ]
    for para in c59:
        add_body(doc, para, first_indent=True)

    # --- Conținut extins (~40 pagini țintă): caracteristici, API, scenarii, glosar ---
    add_extra_from_file(doc)

    page_break(doc)

    # ===================== CAP 6 =====================
    add_center_title(doc, "6. Concluzii")
    conc = [
        "Lucrarea a demonstrat fezabilitatea integrării unui motor de anomalii bazat pe Isolation Forest "
        "împreună cu un pipeline RAG hibrid pentru contextualizarea alertelor Wazuh. Arhitectura "
        "propusă respectă constrângeri operaționale: date reale zgomotoase, necesitatea explicațiilor "
        "și controlul utilizatorului asupra excepțiilor.",

        "Gradul de îndeplinire al obiectivelor este ridicat: implementarea este funcțională, modulară "
        "și include evaluare cantitativă în două faze (baseline vs reantrenat). Îmbunătățirile "
        "observate pe setul mare de antrenare confirmă utilitatea reantrenării periodice.",

        "Dezvoltări viitoare recomandate includ: cross-encoder pentru re-ranking, învățare activă "
        "din feedback-ul analistului, extinderea setului de test cu mai multe scenarii de atac "
        "documentate, și monitorizare automată a deriva scorurilor (drift) cu alertare către echipa "
        "SOC.",

        "Limitări: etichetarea automată nu este perfectă; LLM poate genera erori dacă contextul este "
        "insuficient sau dacă porțile de nivel nu sunt respectate; infrastructura necesită mentenanță "
        "(containere, modele embedding, spațiu disc).",

        "În ansamblu, soluția oferă un echilibru între automatizare și control uman, fiind potrivită "
        "ca prototip extins spre producție în laboratoare de securitate sau ca bază pentru o teză "
        "de master pe tema învățării continue în SOC.",
    ]
    for para in conc:
        add_body(doc, para, first_indent=True)

    page_break(doc)

    # ===================== Bibliografie =====================
    add_center_title(doc, "Bibliografie")
    refs = [
        "[1] F. T. Liu, K. M. Ting, Z.-H. Zhou, „Isolation forest”, în Proc. IEEE ICDM, 2008, pp. 413–422.",
        "[2] C. C. Aggarwal, Outlier Analysis, 2nd ed., Springer, 2016.",
        "[3] P. Lewis et al., „Retrieval-augmented generation for knowledge-intensive NLP tasks”, NeurIPS, 2020.",
        "[4] N. Reimers, I. Gurevych, „Sentence-BERT: Sentence embeddings using Siamese BERT-networks”, EMNLP-IJCNLP, 2019.",
        "[5] MITRE ATT&CK, disponibil: https://attack.mitre.org/",
        "[6] MITRE D3FEND, disponibil: https://d3fend.mitre.org/",
        "[7] Documentație Wazuh, disponibil: https://documentation.wazuh.com/",
        "[8] Qdrant Documentation, disponibil: https://qdrant.tech/documentation/",
        "[9] PostgreSQL Documentation, disponibil: https://www.postgresql.org/docs/",
        "[10] scikit-learn: IsolationForest, disponibil: https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.IsolationForest.html",
        "[11] CISA Known Exploited Vulnerabilities Catalog, disponibil: https://www.cisa.gov/known-exploited-vulnerabilities-catalog",
        "[12] SigmaHQ, repository Sigma, disponibil: https://github.com/SigmaHQ/sigma",
        "[13] OWASP Top 10, disponibil: https://owasp.org/www-project-top-ten/",
        "[14] S. Chuvakin, K. Schmidt, Logging and Log Management, Syngress, 2013.",
        "[15] Rapoarte ENISA / DBIR (ed. recentă) pentru contextul amenințărilor — consultați ediția anului curent.",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.3
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    page_break(doc)

    # ===================== Anexe =====================
    add_center_title(doc, "Anexa 1 — Lista figurilor (de completat în Word)")
    for item in (
        "Figura 4.1. Diagrama componentelor sistemului.",
        "Figura 4.2. Fluxul funcțional principal al /api/chat.",
        "Figura 5.1. Matrice de confuzie (opțional, din evaluate_isolation_forest.py --plot).",
    ):
        add_body(doc, item, first_indent=False)
    page_break(doc)

    add_center_title(doc, "Anexa 2 — Fragmente de cod sursă (Courier New 10 bold în lucrarea finală)")
    add_body(
        doc,
        "Inserați aici, din editor, fragmentele relevante (sub 3/4 pagină fiecare) din: "
        "ai_engine/anomaly_detector.py (extragerea caracteristicilor și normalizarea scorului), "
        "train_isolation_forest.py (antrenare pe benign), rag_core/database/qdrant_store.py "
        "(căutare hibridă RRF), backend/server.py (orchestare chat și prag nivel alertă).",
        first_indent=True,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"S-a generat: {OUTPUT}")


if __name__ == "__main__":
    main()
